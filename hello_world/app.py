import json
import os
from docx import Document
import boto3
from io import BytesIO
import google.generativeai as genai
from cuid2 import Cuid
import re

def essayPromp(topic, essay):
  prompt = f"You are an IELTS task 2 writing examninor. Your job is to grade IELTS essay. Consider coherence, cohesion, vocabulary, grammar, and task achievement. Provide specific feedback on each aspect, highlighting strengths and areas for improvement. Additionally, assess the essay's overall structure, clarity of arguments, and use of examples. At the end, provide an improved version of the essay. The topic and essay will be provided below. \n Topic: ${topic}\n Essay: {essay}";
  return prompt



def fileProcessor(text, essay):
    document = Document()
    lines = text.splitlines()
    paragraph = document.add_paragraph(essay)
    for line in lines:
    # Check if line starts with '**' (bold formatting)
        if line.startswith('**'):
            # Remove the starting '**' and ending characters
            bold_text = line[2:-2]
            # Create a new paragraph and make the text bold
            paragraph = document.add_paragraph()
            run = paragraph.add_run(bold_text)
            run.bold = True
        elif line.startswith('*'):
            # Remove the starting '*' and add as a bullet point
            paragraph = document.add_paragraph()
            paragraph.style = document.styles['List Bullet']  # Set bullet point style
            nextText = line[1:].strip()
            if (nextText.startswith("**")):
                pattern = r"\*\*(.*?)\*\*"
                match = re.search(pattern, nextText)


                if match:
                    matched_bold_text = match.group(1);
                    run = paragraph.add_run(matched_bold_text)
                    run.bold = True
                    text_after_bold = nextText[match.end():]
                    run = paragraph.add_run(text_after_bold)
            else:
                paragraph.add_run(nextText)
        else:
            # Add the line as a normal paragraph
            document.add_paragraph(line)
    file_content = BytesIO()
    document.save(file_content)
    file_content.seek(0)
    return file_content

    



def lambda_handler(event, context):

    # Extract topic and essay
    body = event.get('body')
    if not body:
      return {
          'statusCode': 400,
          'body': 'Empty request body'
      }
    try:
      data = json.loads(body)
    except json.JSONDecodeError:
      return {
          'statusCode': 400,
          'body': 'Invalid JSON format'
      }
    topic = data.get('topic')
    essay = data.get('essay')

    # Generate answer from gemini
    GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY')
    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel('gemini-pro')
    prompt = essayPromp(topic, essay)
    ai_reponse = model.generate_content(prompt)

    # File Processing
    file_content = fileProcessor(ai_reponse.text, essay)


    # Upload file to s3 
    CUID_GENERATOR = Cuid(length = 15)
    file_id = CUID_GENERATOR.generate()
    s3_client = boto3.client('s3')
    s3_client.put_object(Body=file_content, Bucket=os.environ.get('BUCKET_NAME_GRADED'), Key=f'{file_id}.docx')
        

    data = {
       'file': file_id
    }
    

    return {
      'statusCode': 200,
      'body': json.dumps(data)
    }


